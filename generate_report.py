"""
generate_report.py
Generates a professional 5-section academic / business report for the
Zentai Networks Capstone project as a .docx Word document.

Run from the project root:
    python3 generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    """Adds a formatted section heading."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)  # Dark navy
    return p

def add_body(doc, text, indent=False):
    """Adds a clean body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Adds a bullet-point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_b = p.add_run(bold_prefix + ": ")
        set_font(run_b, size=11, bold=True)
        run = p.add_run(text)
        set_font(run, size=11)
    else:
        run = p.add_run(text)
        set_font(run, size=11)
    return p

def add_table_row(table, cells, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, size=10, bold=bold)
    return row

def set_table_header_style(row):
    """Makes the header row dark navy with white text."""
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "003366")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)

def add_divider(doc):
    """Adds a thin horizontal rule between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    ppr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003366")
    pb.append(bottom)
    ppr.append(pb)

# ─────────────────────────────────────────────────────────────────────────────
# Main report generation
# ─────────────────────────────────────────────────────────────────────────────

def generate():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Cover / Title Block ───────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Zentai Networks")
    set_font(run, name="Calibri", size=28, bold=True, color=(0x00, 0x33, 0x66))

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Enterprise Business Intelligence & Decision Support System")
    set_font(run, name="Calibri", size=16, italic=True, color=(0x44, 0x44, 0x44))

    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.date.today().strftime("%B %d, %Y")
    run = meta_p.add_run(f"Capstone Project Report  |  {today}")
    set_font(run, size=11, color=(0x66, 0x66, 0x66))

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    # ── Section 1: Executive Summary ─────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)
    add_body(doc,
        "Zentai Networks is a next-generation, production-grade Business Intelligence and "
        "Decision Support System engineered to serve the strategic information needs of "
        "founders, Chief Financial Officers, and senior leadership. The platform consolidates "
        "multi-dimensional financial data — live market prices, revenue fundamentals, "
        "geopolitical supply chain exposure, and real-time news sentiment — into a single, "
        "unified analytical interface."
    )
    add_body(doc,
        "The primary deliverable of the system is the Unified Business Health Score, a "
        "mathematically derived scalar metric that compresses the output of three independent "
        "AI analytical models into a single, instantly readable executive directive. This "
        "addresses the core challenge of modern executive decision-making: information overload "
        "and analytical fragmentation across disparate data sources."
    )
    add_body(doc,
        "The system is built entirely in Python, using Streamlit for the reactive frontend, "
        "Pandas and NumPy for high-performance data processing, SciPy for statistical "
        "computation, and Plotly for WebGL-accelerated interactive visualisations. All market "
        "data is sourced in real time from Yahoo Finance's REST APIs via the yfinance library, "
        "ensuring zero data staleness."
    )
    add_divider(doc)

    # ── Section 2: Problem Statement ─────────────────────────────────────────
    add_heading(doc, "2. Problem Statement & Business Motivation", level=1)
    add_body(doc,
        "The modern CFO or senior executive faces a fundamental information management "
        "problem. Strategic decisions that may deploy millions of dollars of capital are "
        "routinely made by manually synthesising data from Bloomberg terminals, quarterly "
        "earnings reports, supply chain dashboards, and news aggregators. This workflow "
        "introduces three critical failure modes:"
    )
    add_bullet(doc,
        "Data is captured from multiple independent tools with different update frequencies, "
        "creating temporal inconsistencies in the analytical picture.",
        bold_prefix="Data Fragmentation"
    )
    add_bullet(doc,
        "Legacy BI tools are fundamentally reactive — they display what has already happened. "
        "They cannot mathematically project future risk, evaluate intrinsic valuation, or "
        "identify forward-looking market patterns.",
        bold_prefix="Reactive Analysis Paradigm"
    )
    add_bullet(doc,
        "Presenting raw, unformatted datasets to executive stakeholders without automated "
        "synthesis leads to analysis paralysis and significantly increases decision latency "
        "in time-sensitive market environments.",
        bold_prefix="Decision Paralysis"
    )
    add_body(doc,
        "Zentai Networks directly addresses each of these failure modes by centralising all "
        "data pipelines into a single platform and automating the synthesis of complex signals "
        "into clear, actionable, human-readable directives."
    )
    add_divider(doc)

    # ── Section 3: System Architecture & Technical Design ────────────────────
    add_heading(doc, "3. System Architecture & Technical Design", level=1)

    add_heading(doc, "3.1  Architectural Pattern", level=2)
    add_body(doc,
        "The application is implemented as a Modular Monolith. This architectural pattern "
        "was deliberately chosen over a Microservices architecture for the following reason: "
        "the platform's quantitative engines perform heavy Pandas DataFrame operations on "
        "shared in-memory data structures. Decomposing these into independent services would "
        "require serialising large DataFrames into JSON payloads for inter-service "
        "communication, introducing prohibitive network latency that would severely degrade "
        "the real-time user experience."
    )
    add_body(doc,
        "The system is separated into three strictly decoupled layers:"
    )
    add_bullet(doc, "Data Ingestion Layer — REST API calls to Yahoo Finance via yfinance.", bold_prefix="Layer 1")
    add_bullet(doc, "Mathematical Processing Engines — Modular Python functions for each analytical model.", bold_prefix="Layer 2")
    add_bullet(doc, "Frontend UI Renderer — Streamlit components that consume the engine outputs.", bold_prefix="Layer 3")

    add_heading(doc, "3.2  Technology Stack", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["Component", "Technology", "Purpose"]):
        hdr.cells[i].text = h
    set_table_header_style(hdr)

    rows_data = [
        ("Language",          "Python 3.12",                    "Core runtime and OOP engine architecture."),
        ("Frontend",          "Streamlit",                      "React-based UI generated entirely from Python."),
        ("Data Science",      "Pandas / NumPy",                 "Vectorised data operations and linear algebra."),
        ("Statistics",        "SciPy",                          "Pearson Correlation for pattern matching."),
        ("Visualisation",     "Plotly (WebGL)",                  "High-performance interactive charts and 3D globe."),
        ("Market Data",       "yfinance",                       "Live REST API wrapper for Yahoo Finance."),
        ("State Management",  "st.session_state",               "Persistent browser-session variable storage."),
        ("UI Styling",        "Custom CSS Keyframes",           "Glassmorphism animations injected into the DOM."),
    ]
    for row_data in rows_data:
        r = table.add_row()
        for i, text in enumerate(row_data):
            r.cells[i].text = text
            for para in r.cells[i].paragraphs:
                for run in para.runs:
                    set_font(run, size=10)

    doc.add_paragraph()

    add_heading(doc, "3.3  Data Entity Mapping", level=2)
    add_body(doc,
        "The entire application is built around a single Python dictionary, COMPANY_MAP, "
        "which serves as the authoritative registry for all 10 companies tracked by the "
        "platform. Each entry stores the Yahoo Finance ticker symbol, the company's industry "
        "classification, and its brand color. This O(1)-lookup registry ensures that adding "
        "a new company to the platform requires only a single dictionary entry — every "
        "downstream engine immediately picks it up without code duplication."
    )
    add_divider(doc)

    # ── Section 4: Core Analytical Engines ───────────────────────────────────
    add_heading(doc, "4. Core Analytical Engines", level=1)

    add_heading(doc, "4.1  Unified Business Health Score (AI Ensemble Engine)", level=2)
    add_body(doc,
        "The Unified Business Health Score is the central deliverable of the Capstone Rubric. "
        "Rather than relying on a single data signal — which is inherently noisy and "
        "unreliable — the platform runs three independent AI models in parallel and aggregates "
        "their votes into a single Consensus Directive."
    )
    add_bullet(doc, "Calculates the standard deviation of recent closing prices to evaluate market-implied risk.", bold_prefix="Model 1 — Risk AI")
    add_bullet(doc, "Evaluates gross margins and revenue growth trends against hardcoded algorithmic thresholds.", bold_prefix="Model 2 — Financial AI")
    add_bullet(doc, "Uses heuristic keyword matching against live news headlines to generate a sentiment score.", bold_prefix="Model 3 — Strategy AI")
    add_body(doc,
        "Each model votes +1 (Buy), 0 (Hold), or -1 (Sell). The algebraic sum of all three "
        "votes produces the final Consensus: STRONG BUY (+3/+2), BUY (+1), HOLD (0), "
        "SELL (-1), or STRONG SELL (-2/-3). A confidence percentage is calculated as the "
        "absolute decisiveness of the vote divided by the maximum possible score."
    )

    add_heading(doc, "4.2  Discounted Cash Flow (DCF) Valuation Engine", level=2)
    add_body(doc,
        "The DCF Engine implements the gold standard of institutional financial analysis. "
        "The algorithm fetches the company's latest Free Cash Flow (FCF) figure from the "
        "Yahoo Finance fundamentals feed, projects it five years into the future using a "
        "constant growth rate assumption, and discounts the projected cash flows back to "
        "their present value using an estimated Weighted Average Cost of Capital (WACC). "
        "The resulting Intrinsic Value is compared to the current market price to produce a "
        "Margin of Safety percentage, which tells the executive precisely whether the stock "
        "is fundamentally undervalued or overvalued by the market."
    )

    add_heading(doc, "4.3  Portfolio Optimization Engine (Modern Portfolio Theory)", level=2)
    add_body(doc,
        "The portfolio engine implements Harry Markowitz's Nobel Prize-winning Modern "
        "Portfolio Theory (MPT). The algorithm fetches 365 days of daily closing prices for "
        "all 10 assets, calculates their daily percentage returns, and constructs a 10×10 "
        "Covariance Matrix. This matrix captures how every pair of assets moves relative to "
        "each other — the foundational insight of MPT."
    )
    add_body(doc,
        "Optimal weights are calculated using Inverse Volatility Weighting: assets with high "
        "standard deviations (such as Tesla) receive lower capital allocations, while stable "
        "assets (such as JPMorgan Chase) receive heavier allocations. This mathematically "
        "minimises total portfolio risk and maximises the Sharpe Ratio. The Capital Allocator "
        "then multiplies these weights against a user-specified dollar amount to generate a "
        "precise, executable allocation table."
    )

    add_heading(doc, "4.4  Algorithmic Backtesting Engine", level=2)
    add_body(doc,
        "The backtester evaluates three quantitative trading strategies against historical "
        "price data using fully vectorised Pandas operations. The most critical engineering "
        "consideration is look-ahead bias prevention — achieved by applying df['Signal'].shift(1), "
        "which ensures the algorithm only acts on information that was available the prior "
        "trading day, perfectly simulating real-world execution constraints. Supported "
        "strategies include SMA Crossover (trend-following), RSI Mean Reversion (momentum "
        "oscillator), and MACD Momentum (trend acceleration)."
    )

    add_heading(doc, "4.5  Geospatial Supply Chain Engine (Haversine Risk Collider)", level=2)
    add_body(doc,
        "This engine maps each company's real-world physical infrastructure — manufacturing "
        "plants, data centers, and distribution hubs — onto a 3D interactive Plotly globe. "
        "A mathematical risk-collision algorithm then calculates whether any infrastructure "
        "node is physically within the blast radius of an active geopolitical danger zone "
        "(e.g., the Taiwan Strait, the Red Sea). Because the Earth is a sphere, standard "
        "Euclidean distance fails; the engine implements the Haversine Formula to calculate "
        "the precise great-circle distance between two latitude/longitude coordinates."
    )

    add_heading(doc, "4.6  Historical Pattern Recognition Engine", level=2)
    add_body(doc,
        "This engine transitions the platform from static reporting to predictive analytics. "
        "It scans five years of historical closing price data and uses Pearson Correlation "
        "(via scipy.stats.pearsonr) to find the exact 30-day period in history that "
        "mathematically mirrors the current 30-day trend. Both price arrays are first "
        "normalised using Z-score standardisation to ensure fair comparison across different "
        "price levels. The Predictive Shadow then explicitly calculates the return "
        "achieved in the 30 days immediately following the historical match, providing a "
        "quantitative statistical baseline projection for the coming month."
    )
    add_divider(doc)

    # ── Section 5: Capstone Alignment, Limitations & Roadmap ─────────────────
    add_heading(doc, "5. Capstone Alignment, Limitations & Future Roadmap", level=1)

    add_heading(doc, "5.1  Alignment with Capstone Rubric", level=2)
    add_body(doc,
        "Every module of the Zentai platform maps directly to a specific objective in the "
        "Capstone Rubric:"
    )

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0]
    for i, h in enumerate(["Capstone Objective", "Platform Implementation"]):
        hdr2.cells[i].text = h
    set_table_header_style(hdr2)

    alignment_data = [
        ("Unified Business Health Score",       "AI Ensemble Engine — three-model vote producing a single Consensus Directive."),
        ("Financial Performance Aggregation",   "DCF Valuation Engine and live KPI fetching (revenue, margins, growth)."),
        ("Operational Efficiency Metrics",      "Geospatial Supply Chain Heatmap with Haversine risk collision detection."),
        ("Forward-Looking Risk Insights",       "Historical Pattern Recognition Engine with Predictive Shadow projections."),
        ("Proactive Financial Planning",        "Portfolio Optimizer and Algorithmic Backtester with Alpha measurement."),
        ("Strategic Decision Support",          "Executive Briefing Room AI chatbot and Sector War-Room competitor analysis."),
    ]
    for row_data in alignment_data:
        r = table2.add_row()
        for i, text in enumerate(row_data):
            r.cells[i].text = text
            for para in r.cells[i].paragraphs:
                for run in para.runs:
                    set_font(run, size=10)

    doc.add_paragraph()

    add_heading(doc, "5.2  Known Limitations", level=2)
    add_bullet(doc,
        "All market data is sourced from Yahoo Finance, which is subject to rate limiting and "
        "occasional data gaps. During off-market hours, some metrics (e.g., live options volume) "
        "will return zero or stale values.",
        bold_prefix="Data Source Dependency"
    )
    add_bullet(doc,
        "The geographic capital flow data in the Market Intelligence module is algorithmically "
        "simulated based on real volume sizes. Stock exchanges do not legally disclose the "
        "physical origin of individual orders.",
        bold_prefix="Simulated Geographic Data"
    )
    add_bullet(doc,
        "Historical pattern recognition relies on Pearson Correlation, a linear similarity "
        "measure. Non-linear market patterns (e.g., volatility clustering) may be missed.",
        bold_prefix="Linear Correlation Assumption"
    )

    add_heading(doc, "5.3  Future Architecture Roadmap", level=2)
    add_bullet(doc,
        "Replace the yfinance REST polling mechanism with authenticated WebSocket connections "
        "to major exchange feeds for millisecond-level tick data.",
        bold_prefix="Real-Time Data Layer"
    )
    add_bullet(doc,
        "Integrate PostgreSQL to persist custom portfolio configurations, historical backtest "
        "results, and executive notes across sessions.",
        bold_prefix="Persistent Data Storage"
    )
    add_bullet(doc,
        "Connect the Capital Allocator engine directly to institutional brokerage APIs "
        "(e.g., Alpaca Markets) to enable one-click, algorithm-driven trade execution.",
        bold_prefix="Automated Execution Layer"
    )
    add_bullet(doc,
        "Package the backend engines as independent FastAPI microservices and containerise "
        "them via Docker for scalable cloud deployment on AWS ECS or Google Cloud Run.",
        bold_prefix="Cloud-Native Containerisation"
    )

    add_divider(doc)

    # ── Footer paragraph ──────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        f"Zentai Networks — Capstone Project Report  |  {today}  |  Confidential"
    )
    set_font(run, size=9, italic=True, color=(0x88, 0x88, 0x88))

    # ── Save ──────────────────────────────────────────────────────────────────
    output = "Zentai_Networks_Capstone_Report.docx"
    doc.save(output)
    print(f"Report saved → {output}")


if __name__ == "__main__":
    generate()
